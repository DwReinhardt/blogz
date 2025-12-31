
"""
Collate Resume Bullets (Role-separated Experience) -> Excel (.xlsx) with SourceFile

Purpose:
- Parse multiple .docx resumes and aggregate bullet points into an Excel workbook.
- Separate bullets under Experience/Professional Experience by detected "role" headers (e.g., Title — Company — Dates).
- Include bullets from other configured sections on their own sheets.
- Add a SourceFile column to every row to trace bullets to the original resume.
- Deduplicate bullets (case-insensitive).
- Prompt for input/output paths if not provided; optional GUI pickers via --gui.

Sheets created:
- 'Experience (by Role)'               : columns [Role, Bullet, SourceFile]
- 'Professional Experience (by Role)'  : columns [Role, Bullet, SourceFile]
- Other sections (one sheet each)      : columns [Bullet, SourceFile]

Usage:
    # Console prompts
    python collate_resume_bullets_roles_to_excel_prompt.py

    # GUI pickers
    python collate_resume_bullets_roles_to_excel_prompt.py --gui

    # CLI arguments (no prompts)
    python collate_resume_bullets_roles_to_excel_prompt.py "<input_folder>" "<output_xlsx>"

Dependencies:
    pip install python-docx pandas openpyxl

Author:
- Adapted for Duane Reinhardt with expanded sections, SourceFile column, and documentation.
"""

import os
import re
import sys
from collections import defaultdict
from docx import Document
import pandas as pd

# =============================================================================
# 1) CONFIGURATION & SECTION CANONICALIZATION
# =============================================================================
# --- Sections to include (canonical set) ---
SECTIONS_TO_INCLUDE = {
    "Professional Profile",
    "Technical Skills",
    "Professional Experience",
    "Additional Skills & Qualifications",
    "Experience",
    "Core Competencies",
    "Skills",
    "Education",
    "Certifications",
    "Summary",
}

# Preferences
DEDUPLICATE_BULLETS = True            # remove duplicate bullets (case-insensitive)
BULLET_CHARS = {"•", "-", "–", "*", "·"}  # common bullet glyphs

# Canonical mapping: normalize common variants to the canonical set above
CANONICAL_MAP = {
    # Experience family
    "Work Experience": "Experience",
    "Employment": "Experience",
    "Professional Experience": "Professional Experience",
    # Summary/Profile
    "Professional Summary": "Summary",
    "Objective": "Summary",
    "Profile": "Professional Profile",
    # Skills family
    "Technical Skills": "Technical Skills",
    "Core Competency": "Core Competencies",
    "Competencies": "Core Competencies",
    "Additional Skills": "Additional Skills & Qualifications",
    "Additional Qualifications": "Additional Skills & Qualifications",
    # Certifications
    "Certs": "Certifications",
}

# Heading detection regexes (case-insensitive)
HEADING_REGEXES = [
    r"^\s*(professional profile|profile)\s*:?\s*$",
    r"^\s*(technical skills)\s*:?\s*$",
    r"^\s*(professional experience|experience|work experience|employment)\s*:?\s*$",
    r"^\s*(additional skills & qualifications|additional skills|additional qualifications)\s*:?\s*$",
    r"^\s*(core competencies|competencies|core competency)\s*:?\s*$",
    r"^\s*(skills)\s*:?\s*$",
    r"^\s*(education)\s*:?\s*$",
    r"^\s*(certifications|certs)\s*:?\s*$",
    r"^\s*(summary|objective|professional summary)\s*:?\s*$",
]
HEADING_COMPILED = [re.compile(rx, re.IGNORECASE) for rx in HEADING_REGEXES]

# Date patterns commonly found in role headers (e.g., "Jan 2020 – Present", "2019–2024", "1/2020–12/2023")
DATE_PATTERNS = [
    r"\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)\s+\d{4}\b",
    r"\b\d{4}\b",
    r"\b\d{1,2}/\d{4}\b",
]
DATE_RANGE = re.compile(
    r"(" + "|".join(DATE_PATTERNS) + r")\s*(?:[-–—]\s*|to\s+)\s*(?:Present|Current|Now|" + "|".join(DATE_PATTERNS) + r")",
    re.IGNORECASE
)

# =============================================================================
# 2) INPUT HELPERS: CONSOLE PROMPTS & OPTIONAL GUI
# =============================================================================
def prompt_for_folder(prompt_text: str) -> str:
    """
    Console prompt to capture a valid folder path.
    Ensures the folder exists before returning.
    """
    while True:
        path = input(f"{prompt_text}\n> ").strip().strip('"')
        if not path:
            print("Please enter a non-empty path.")
            continue
        if not os.path.isdir(path):
            print(f"Path does not exist or is not a folder: {path}")
            continue
        return path

def prompt_for_output_xlsx(prompt_text: str) -> str:
    """
    Console prompt to capture a valid .xlsx output path.
    Ensures .xlsx extension and creates parent directory if needed.
    """
    while True:
        path = input(f"{prompt_text}\n> ").strip().strip('"')
        if not path:
            print("Please enter a non-empty path.")
            continue
        if not path.lower().endswith(".xlsx"):
            print("Output must end with .xlsx (e.g., C:\\...\\Collated Bullets - Resumes.xlsx)")
            continue
        parent = os.path.dirname(path) or "."
        try:
            os.makedirs(parent, exist_ok=True)
        except Exception as e:
            print(f"Cannot create/access parent folder '{parent}': {e}")
            continue
        return path

def pick_paths_gui() -> tuple[str, str]:
    """
    Simple GUI pickers using tkinter:
    - askdirectory for input folder
    - asksaveasfilename for output .xlsx
    Falls back to console prompts if tkinter isn't available.
    """
    try:
        import tkinter as tk
        from tkinter import filedialog
        root = tk.Tk()
        root.withdraw()

        input_folder = filedialog.askdirectory(title="Select folder containing .docx resumes")
        if not input_folder:
            print("No input folder selected. Falling back to console prompt.")
            input_folder = prompt_for_folder("Enter the input folder path (OneDrive-synced folder with .docx resumes):")

        output_xlsx = filedialog.asksaveasfilename(
            title="Choose output .xlsx",
            defaultextension=".xlsx",
            filetypes=[("Excel Workbook", "*.xlsx")],
            initialfile="Collated Bullets - Resumes.xlsx"
        )
        if not output_xlsx:
            print("No output file selected. Falling back to console prompt.")
            output_xlsx = prompt_for_output_xlsx("Enter the full output .xlsx path:")

        return input_folder, output_xlsx
    except Exception:
        print("GUI pickers unavailable. Using console prompts.")
        input_folder = prompt_for_folder("Enter the input folder path (OneDrive-synced folder with .docx resumes):")
        output_xlsx = prompt_for_output_xlsx("Enter the full output .xlsx path:")
        return input_folder, output_xlsx

# =============================================================================
# 3) PARSING HELPERS: SECTION HEADINGS, BULLET LINES, ROLE HEADERS
# =============================================================================
def canonical_section(name: str) -> str:
    """
    Normalize a section name to the canonical set in SECTIONS_TO_INCLUDE.
    Returns Title-Case fallback for unknown names; 'Uncategorized' when missing.
    """
    n = (name or "").strip()
    if not n:
        return "Uncategorized"
    for k, v in CANONICAL_MAP.items():
        if k.lower() == n.lower():
            return v
    for s in SECTIONS_TO_INCLUDE:
        if s.lower() == n.lower():
            return s
    return n.title()

def is_heading_paragraph(p) -> str | None:
    """
    Decide if 'p' is a section heading using:
    - Word heading styles (Heading 1..3),
    - ALL CAPS short line heuristic,
    - Regex matching for common headings.
    Returns canonical section name or None.
    """
    text = (p.text or "").strip()
    style_name = (getattr(p.style, "name", "") or "").lower()

    if "heading" in style_name:
        return canonical_section(text)
    if text and len(text) <= 60 and text.isupper():
        return canonical_section(text.title())
    for rx in HEADING_COMPILED:
        if rx.match(text):
            return canonical_section(text.title())
    return None

def is_bullet_paragraph(p) -> bool:
    """
    Detect bullet/numbered list lines via:
    - Paragraph style name containing 'list' or 'bullet'
    - Leading bullet glyphs (• - – * ·)
    - Numbered patterns like '1. ', 'a) ', 'A. '
    """
    style_name = (getattr(p.style, "name", "") or "").lower()
    if "list" in style_name or "bullet" in style_name:
        return True

    text = (p.text or "").strip()
    if not text:
        return False

    if text[0] in BULLET_CHARS:
        return True
    if re.match(r"^(\d+\.|\d+\)|[a-zA-Z]\)|[a-zA-Z]\.)\s+", text):
        return True
    return False

def clean_bullet_text(text: str) -> str:
    """
    Remove bullet/number markers and normalize whitespace.
    """
    t = (text or "").strip()
    t = re.sub(r"^[\u2022\-\–\*\·]+\s+", "", t)  # bullet glyphs
    t = re.sub(r"^(\d+\.|\d+\)|[a-zA-Z]\)|[a-zA-Z]\.)\s+", "", t)  # numbering tokens
    t = re.sub(r"\s+", " ", t)
    return t.strip()

def paragraph_has_bold_run(p) -> bool:
    """
    True if any run in paragraph is bold—often indicates role headers.
    """
    try:
        for run in p.runs:
            if run.bold:
                return True
    except Exception:
        pass
    return False

def looks_like_role_header(text: str) -> bool:
    """
    Heuristics to detect role headers within Experience-like sections:
    - Presence of a date range (e.g., 'Jan 2020 – Present', '2019–2024')
    - Common separators for title/company (' - ', ' — ', ' | ')
    - Title-case density with reasonable length
    """
    t = (text or "").strip()
    if not t:
        return False
    if DATE_RANGE.search(t):
        return True
    if " - " in t or " — " in t or " | " in t:
        return True
    words = [w for w in re.split(r"\s+", t) if w]
    titleish = sum(1 for w in words if w[:1].isupper())
    if len(words) >= 2 and titleish / max(len(words), 1) > 0.6 and len(t) <= 120:
        return True
    return False

# =============================================================================
# 4) EXTRACTION PER RESUME: EXPERIENCE ROLES & OTHER SECTIONS
# =============================================================================
def extract_roles_and_bullets_from_experience(paragraphs) -> dict[str, list[str]]:
    """
    Given a list of paragraphs belonging to an Experience-like section, detect role headers
    and collect bullets under each role until the next role header.

    Returns: {role_header_text -> [bullets]}
    """
    roles: dict[str, list[str]] = defaultdict(list)
    current_role = None

    for p in paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue

        is_role_header = looks_like_role_header(text) or paragraph_has_bold_run(p)
        if is_role_header and not is_bullet_paragraph(p):
            current_role = text
            roles[current_role] = roles.get(current_role, [])
            continue

        if is_bullet_paragraph(p):
            bullet = clean_bullet_text(text)
            if bullet:
                if not current_role:
                    current_role = "Role: Unknown"
                    roles[current_role] = roles.get(current_role, [])
                roles[current_role].append(bullet)

    return roles

def extract_bullets_with_roles_from_docx(path: str):
    """
    Parse a single .docx resume and return:
      - exp_roles     : {role_header -> [bullets]}  # for 'Experience'
      - profexp_roles : {role_header -> [bullets]}  # for 'Professional Experience'
      - other_sections: {section -> [bullets]}      # all other sections
    """
    doc = Document(path)
    current_section = None
    section_to_paragraphs: dict[str, list] = defaultdict(list)

    # Gather paragraphs by active section (flow through document order)
    for p in doc.paragraphs:
        heading = is_heading_paragraph(p)
        if heading:
            current_section = heading
            continue
        sec = canonical_section(current_section or "Uncategorized")
        section_to_paragraphs[sec].append(p)

    # Include paragraphs from tables (common on resumes)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    heading = is_heading_paragraph(p)
                    if heading:
                        current_section = heading
                        continue
                    sec = canonical_section(current_section or "Uncategorized")
                    section_to_paragraphs[sec].append(p)

    exp_roles: dict[str, list[str]] = defaultdict(list)
    profexp_roles: dict[str, list[str]] = defaultdict(list)
    other_sections: dict[str, list[str]] = defaultdict(list)

    for sec, plist in section_to_paragraphs.items():
        sec_lower = sec.lower()
        if sec_lower == "experience":
            roles_map = extract_roles_and_bullets_from_experience(plist)
            for role, bullets in roles_map.items():
                if bullets:
                    exp_roles[role].extend(bullets)
        elif sec_lower == "professional experience":
            roles_map = extract_roles_and_bullets_from_experience(plist)
            for role, bullets in roles_map.items():
                if bullets:
                    profexp_roles[role].extend(bullets)
        else:
            for p in plist:
                if is_bullet_paragraph(p):
                    bullet = clean_bullet_text(p.text)
                    if bullet:
                        other_sections[sec].append(bullet)

    # Per-document dedup
    if DEDUPLICATE_BULLETS:
        for role, items in list(exp_roles.items()):
            seen = set(); unique = []
            for b in items:
                k = b.strip().lower()
                if k not in seen:
                    seen.add(k); unique.append(b)
            exp_roles[role] = unique

        for role, items in list(profexp_roles.items()):
            seen = set(); unique = []
            for b in items:
                k = b.strip().lower()
                if k not in seen:
                    seen.add(k); unique.append(b)
            profexp_roles[role] = unique

        for sec, items in list(other_sections.items()):
            seen = set(); unique = []
            for b in items:
                k = b.strip().lower()
                if k not in seen:
                    seen.add(k); unique.append(b)
            other_sections[sec] = unique

    return exp_roles, profexp_roles, other_sections

# =============================================================================
# 5) FILE DISCOVERY
# =============================================================================
def collect_docx_files(root_folder: str) -> list[str]:
    """
    Walk the root folder (including subfolders) and return all .docx file paths.
    """
    docx_files = []
    for dirpath, _, filenames in os.walk(root_folder):
        for f in filenames:
            if f.lower().endswith(".docx"):
                docx_files.append(os.path.join(dirpath, f))
    return docx_files

# =============================================================================
# 6) MAIN: ARGUMENTS/PROMPTS, AGGREGATION, DEDUP, EXCEL WRITE
# =============================================================================
def main():
    """
    Orchestrates the workflow:
      - Read CLI args or prompt/GUI for input/output paths.
      - Validate folder and collect .docx files.
      - Parse each resume; merge Experience/Professional Experience roles and other sections globally.
      - Deduplicate globally per role and per section.
      - Write a structured .xlsx with multiple sheets, each row including SourceFile.
    """
    # Parse arguments / interactive selection
    use_gui = False
    args = [a for a in sys.argv[1:] if a]
    if "--gui" in args:
        use_gui = True
        args = [a for a in args if a != "--gui"]

    if len(args) >= 2:
        input_folder = args[0].strip().strip('"')
        output_xlsx = args[1].strip().strip('"')
    elif use_gui:
        input_folder, output_xlsx = pick_paths_gui()
    else:
        print("No CLI arguments provided. Interactive mode:")
        input_folder = prompt_for_folder("Enter the input folder path (OneDrive-synced folder with .docx resumes):")
        output_xlsx = prompt_for_output_xlsx("Enter the full output .xlsx path:")

    # Validate inputs & find files
    if not os.path.isdir(input_folder):
        print(f"ERROR: Input folder does not exist: {input_folder}")
        return

    files = collect_docx_files(input_folder)
    if not files:
        print(f"ERROR: No .docx files found in (or under) folder: {input_folder}")
        return

    # Global aggregations with SourceFile tracking
    global_exp_roles: dict[str, list[tuple[str, str]]] = defaultdict(list)       # {Role: [(Bullet, SourceFile), ...]}
    global_profexp_roles: dict[str, list[tuple[str, str]]] = defaultdict(list)   # {Role: [(Bullet, SourceFile), ...]}
    global_sections: dict[str, list[tuple[str, str]]] = defaultdict(list)        # {Section: [(Bullet, SourceFile), ...]}

    for path in files:
        source_file = os.path.basename(path)
        try:
            exp_roles, profexp_roles, other_secs = extract_bullets_with_roles_from_docx(path)
        except Exception as e:
            print(f"WARN: Skipping '{source_file}' due to read error: {e}")
            continue

        for role, bullets in exp_roles.items():
            for b in bullets:
                global_exp_roles[role].append((b, source_file))

        for role, bullets in profexp_roles.items():
            for b in bullets:
                global_profexp_roles[role].append((b, source_file))

        for sec, bullets in other_secs.items():
            canon = canonical_section(sec)
            if canon != "Uncategorized":
                if not any(canon.lower() == s.lower() for s in SECTIONS_TO_INCLUDE):
                    continue
            for b in bullets:
                global_sections[canon].append((b, source_file))

    # Global dedup per role and per section (bullet text only; keep first SourceFile seen)
    if DEDUPLICATE_BULLETS:
        for role, items in list(global_exp_roles.items()):
            seen = set(); unique = []
            for bullet, src in items:
                key = bullet.strip().lower()
                if key not in seen:
                    seen.add(key); unique.append((bullet, src))
            global_exp_roles[role] = unique

        for role, items in list(global_profexp_roles.items()):
            seen = set(); unique = []
            for bullet, src in items:
                key = bullet.strip().lower()
                if key not in seen:
                    seen.add(key); unique.append((bullet, src))
            global_profexp_roles[role] = unique

        for sec, items in list(global_sections.items()):
            seen = set(); unique = []
            for bullet, src in items:
                key = bullet.strip().lower()
                if key not in seen:
                    seen.add(key); unique.append((bullet, src))
            global_sections[sec] = unique

    # Build DataFrames for Excel (all include SourceFile)
    sheets = {}

    # Experience (by Role)
    if global_exp_roles:
        rows = []
        for role, pairs in sorted(global_exp_roles.items(), key=lambda kv: kv[0].lower()):
            for bullet, src in pairs:
                rows.append({"Role": role, "Bullet": bullet, "SourceFile": src})
        sheets["Experience (by Role)"] = pd.DataFrame(rows, columns=["Role", "Bullet", "SourceFile"])

    # Professional Experience (by Role)
    if global_profexp_roles:
        rows = []
        for role, pairs in sorted(global_profexp_roles.items(), key=lambda kv: kv[0].lower()):
            for bullet, src in pairs:
                rows.append({"Role": role, "Bullet": bullet, "SourceFile": src})
        sheets["Professional Experience (by Role)"] = pd.DataFrame(rows, columns=["Role", "Bullet", "SourceFile"])

    # Other sections (Bullet + SourceFile)
    preferred_order = [
        "Professional Profile",
        "Technical Skills",
        "Additional Skills & Qualifications",
        "Core Competencies",
        "Skills",
        "Education",
        "Certifications",
        "Summary",
        "Uncategorized",
    ]
    for sec in preferred_order:
        if sec in global_sections and global_sections[sec]:
            rows = [{"Bullet": bullet, "SourceFile": src} for bullet, src in global_sections[sec]]
            sheets[sec] = pd.DataFrame(rows, columns=["Bullet", "SourceFile"])

    if not sheets:
        print("INFO: No bullets detected.")
        return

    # Write Excel with openpyxl engine (sheet names max length 31)
    try:
        with pd.ExcelWriter(output_xlsx, engine="openpyxl") as writer:
            for sheet_name, df in sheets.items():
                safe_name = sheet_name[:31]
                df.to_excel(writer, sheet_name=safe_name, index=False)
        print(f"Done. Wrote: {output_xlsx}")
    except Exception as e:
        print(f"ERROR: Failed to write Excel: {e}")

if __name__ == "__main__":
    main()
