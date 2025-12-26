# This python script processes multiple .docx resumes in a specified folder,
# extracting bullet points from each resume and collating them into a single document.
# It organizes bullets under their respective sections (like Experience, Skills, Education),
# and within Experience, it further categorizes bullets under individual roles.
# Usage: python collate_resume_bullets.py
# code written by Copilot (GPT-4) and refined by Duane Reinhardt
import os
import re
import sys
from collections import defaultdict
from docx import Document
from docx.shared import Pt

# === Default config (CLI args can override paths) ===
DEFAULT_INPUT_FOLDER = r"c:\Users\dwrei\OneDrive - Reinhardt Family\Resume\Submitted Resumes"
DEFAULT_OUTPUT_DOCX = r"c:\Users\dwrei\OneDrive - Reinhardt Family\Resume\Submitted Resumes\Collated Bullets - Resumes.docx"

SECTIONS_TO_INCLUDE = {
    "Experience",
    "Core Competencies",
    "Skills",
    "Education",
    "Certifications",
    "Summary",
}

DEDUPLICATE_BULLETS = True
TAG_WITH_FILENAME = False  # per your preference
BULLET_CHARS = {"•", "-", "–", "*", "·"}

# Canonical mapping for section names
CANONICAL_MAP = {
    "Work Experience": "Experience",
    "Employment": "Experience",
    "Professional Summary": "Summary",
    "Objective": "Summary",
    "Technical Skills": "Skills",
    "Core Competency": "Core Competencies",
    "Competencies": "Core Competencies",
    "Certs": "Certifications",
}

# Regexes
HEADING_REGEXES = [
    r"^\s*(experience|work experience|employment)\s*:?\s*$",
    r"^\s*(core competencies|competencies|core competency)\s*:?\s*$",
    r"^\s*(skills|technical skills)\s*:?\s*$",
    r"^\s*(education)\s*:?\s*$",
    r"^\s*(certifications|certs)\s*:?\s*$",
    r"^\s*(summary|objective|professional summary)\s*:?\s*$",
]
HEADING_COMPILED = [re.compile(rx, re.IGNORECASE) for rx in HEADING_REGEXES]

# Date patterns commonly found in roles
DATE_PATTERNS = [
    r"\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)\s+\d{4}\b",
    r"\b\d{4}\b",
    r"\b\d{1,2}/\d{4}\b",
]
DATE_RANGE = re.compile(
    r"(" + "|".join(DATE_PATTERNS) + r")\s*(?:[-–—]\s*|to\s+)\s*(?:Present|Current|Now|" + "|".join(DATE_PATTERNS) + r")",
    re.IGNORECASE
)

def canonical_section(name: str) -> str:
    n = (name or "").strip()
    if not n:
        return "Uncategorized"
    for k, v in CANONICAL_MAP.items():
        if k.lower() == n.lower():
            return v
    for s in SECTIONS_TO_INCLUDE:
        if s.lower() == n.lower():
            return s
    return n.title()

def is_heading_paragraph(p) -> str | None:
    text = (p.text or "").strip()
    style_name = (getattr(p.style, "name", "") or "").lower()
    if "heading" in style_name:
        return canonical_section(text)
    if text and len(text) <= 60 and text.isupper():
        return canonical_section(text.title())
    for rx in HEADING_COMPILED:
        if rx.match(text):
            return canonical_section(text.title())
    return None

def is_bullet_paragraph(p) -> bool:
    style_name = (getattr(p.style, "name", "") or "").lower()
    if "list" in style_name or "bullet" in style_name:
        return True
    text = (p.text or "").strip()
    if not text:
        return False
    if text[0] in BULLET_CHARS:
        return True
    if re.match(r"^(\d+\.|\d+\)|[a-zA-Z]\)|[a-zA-Z]\.)\s+", text):
        return True
    return False

def clean_bullet_text(text: str) -> str:
    t = (text or "").strip()
    t = re.sub(r"^[\u2022\-\–\*\·]+\s+", "", t)  # remove bullet glyphs
    t = re.sub(r"^(\d+\.|\d+\)|[a-zA-Z]\)|[a-zA-Z]\.)\s+", "", t)  # numbering markers
    t = re.sub(r"\s+", " ", t)
    return t.strip()

def paragraph_has_bold_run(p) -> bool:
    try:
        for run in p.runs:
            if run.bold:
                return True
    except Exception:
        pass
    return False

def looks_like_role_header(text: str) -> bool:
    """
    Heuristics to detect role headers under Experience.
    - Contains a date range (e.g., 2019–2024, Jan 2020 – Present)
    - OR contains separators typical of title/company (" - ", " — ", " | ")
    - Title-case majority words and limited punctuation
    """
    t = (text or "").strip()
    if not t:
        return False
    if DATE_RANGE.search(t):
        return True
    if " - " in t or " — " in t or " | " in t:
        # Often formatted like "Senior Engineer - Acme Corp | 2019–2024"
        return True
    # Title Case heuristic
    words = [w for w in re.split(r"\s+", t) if w]
    titleish = sum(1 for w in words if w[:1].isupper())
    if len(words) >= 2 and titleish / max(len(words), 1) > 0.6 and len(t) <= 120:
        return True
    return False

def extract_roles_and_bullets_from_experience(paragraphs) -> dict[str, list[str]]:
    """
    Walk paragraphs in the Experience section, detect role headers and collect bullets under each role.
    Returns: {role_header_text -> [bullets]}
    """
    roles: dict[str, list[str]] = defaultdict(list)
    current_role = None

    for p in paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue

        # Role header detection
        is_role_header = looks_like_role_header(text) or paragraph_has_bold_run(p)
        if is_role_header and not is_bullet_paragraph(p):
            current_role = text
            # ensure key exists
            roles[current_role] = roles.get(current_role, [])
            continue

        # Bullets assigned to current_role
        if is_bullet_paragraph(p):
            bullet = clean_bullet_text(text)
            if bullet:
                if not current_role:
                    # If we haven't seen a role header yet, bucket under 'Role: Unknown'
                    current_role = "Role: Unknown"
                    roles[current_role] = roles.get(current_role, [])
                roles[current_role].append(bullet)

    return roles

def extract_bullets_with_roles_from_docx(path: str) -> tuple[dict[str, list[str]], dict[str, list[str]]]:
    """
    Returns:
      experience_roles -> {role_header -> [bullets]}
      other_sections   -> {section -> [bullets]}
    """
    doc = Document(path)

    # First, build a linear stream of paragraphs with section tracking
    current_section = None
    section_to_paragraphs: dict[str, list] = defaultdict(list)

    # paragraphs
    for p in doc.paragraphs:
        heading = is_heading_paragraph(p)
        if heading:
            current_section = heading
            continue
        sec = canonical_section(current_section or "Uncategorized")
        section_to_paragraphs[sec].append(p)

    # tables (include table content in the same stream)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    heading = is_heading_paragraph(p)
                    if heading:
                        current_section = heading
                        continue
                    sec = canonical_section(current_section or "Uncategorized")
                    section_to_paragraphs[sec].append(p)

    # Now parse Experience specially to extract roles
    experience_roles: dict[str, list[str]] = defaultdict(list)
    other_sections: dict[str, list[str]] = defaultdict(list)

    for sec, plist in section_to_paragraphs.items():
        if sec.lower() == "experience":
            roles_map = extract_roles_and_bullets_from_experience(plist)
            # deduplicate within each role
            for role, bullets in roles_map.items():
                if bullets:
                    if DEDUPLICATE_BULLETS:
                        seen = set()
                        unique = []
                        for b in bullets:
                            key = b.strip().lower()
                            if key not in seen:
                                seen.add(key)
                                unique.append(b)
                        bullets = unique
                    experience_roles[role].extend(bullets)
        else:
            # normal section bullets
            for p in plist:
                if is_bullet_paragraph(p):
                    bullet = clean_bullet_text(p.text)
                    if bullet:
                        other_sections[sec].append(bullet)

    # dedup for other sections
    if DEDUPLICATE_BULLETS:
        for sec, items in list(other_sections.items()):
            seen = set()
            unique = []
            for b in items:
                key = b.strip().lower()
                if key not in seen:
                    seen.add(key)
                    unique.append(b)
            other_sections[sec] = unique

    return experience_roles, other_sections

def write_collated_doc_by_roles(experience_roles: dict[str, list[str]],
                                other_sections: dict[str, list[str]],
                                output_path: str):
    out = Document()
    normal = out.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    out.add_heading('Collated Bullets from Resumes (Role-separated Experience)', level=0)

    # Experience first (roles)
    if experience_roles:
        out.add_heading('Experience', level=1)
        # Sort roles to keep a stable order (alphabetical)
        for role in sorted(experience_roles.keys(), key=lambda s: s.lower()):
            out.add_heading(role, level=2)
            for bullet in experience_roles[role]:
                out.add_paragraph(bullet, style='List Bullet')

    # Then other sections in your specified order
    order = ["Core Competencies", "Skills", "Education", "Certifications", "Summary", "Uncategorized"]
    for sec in order:
        if sec in other_sections and other_sections[sec]:
            out.add_heading(sec, level=1)
            for bullet in other_sections[sec]:
                out.add_paragraph(bullet, style='List Bullet')

    out.save(output_path)

def collect_docx_files(root_folder: str) -> list[str]:
    docx_files = []
    for dirpath, _, filenames in os.walk(root_folder):
        for f in filenames:
            if f.lower().endswith(".docx"):
                docx_files.append(os.path.join(dirpath, f))
    return docx_files

def main():
    input_folder = DEFAULT_INPUT_FOLDER
    output_docx = DEFAULT_OUTPUT_DOCX

    if len(sys.argv) >= 2 and sys.argv[1].strip():
        input_folder = sys.argv[1]
    if len(sys.argv) >= 3 and sys.argv[2].strip():
        output_docx = sys.argv[2]

    if not os.path.isdir(input_folder):
        print(f"ERROR: Input folder does not exist: {input_folder}")
        return

    files = collect_docx_files(input_folder)
    if not files:
        print(f"ERROR: No .docx files found in (or under) folder: {input_folder}")
        return

    # Aggregate across all resumes
    global_experience_roles: dict[str, list[str]] = defaultdict(list)
    global_other_sections: dict[str, list[str]] = defaultdict(list)

    for path in files:
        fname = os.path.basename(path)
        try:
            exp_roles, other_secs = extract_bullets_with_roles_from_docx(path)
        except Exception as e:
            print(f"WARN: Skipping '{fname}' due to read error: {e}")
            continue

        # Merge Experience roles
        for role, bullets in exp_roles.items():
            if bullets:
                global_experience_roles[role].extend(bullets)

        # Merge other sections (filter to requested ones + Uncategorized)
        for sec, bullets in other_secs.items():
            canon = canonical_section(sec)
            if canon != "Uncategorized":
                if not any(canon.lower() == s.lower() for s in SECTIONS_TO_INCLUDE):
                    continue
            if bullets:
                global_other_sections[canon].extend(bullets)

    # Global dedup per role and per section
    if DEDUPLICATE_BULLETS:
        for role, items in list(global_experience_roles.items()):
            seen = set()
            unique = []
            for b in items:
                key = b.strip().lower()
                if key not in seen:
                    seen.add(key)
                    unique.append(b)
            global_experience_roles[role] = unique

        for sec, items in list(global_other_sections.items()):
            seen = set()
            unique = []
            for b in items:
                key = b.strip().lower()
                if key not in seen:
                    seen.add(key)
                    unique.append(b)
            global_other_sections[sec] = unique

    non_empty_roles = {k: v for k, v in global_experience_roles.items() if v}
    non_empty_sections = {k: v for k, v in global_other_sections.items() if v}

    if not non_empty_roles and not non_empty_sections:
        print("INFO: No bullets detected.")
        return

    os.makedirs(os.path.dirname(output_docx), exist_ok=True)
    write_collated_doc_by_roles(non_empty_roles, non_empty_sections, output_docx)
    print(f"Done. Wrote: {output_docx}")

if __name__ == "__main__":
    main()